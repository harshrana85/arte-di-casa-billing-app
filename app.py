
import streamlit as st
import pandas as pd
import json, uuid, base64
from pathlib import Path
from datetime import date, datetime
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
ASSETS_DIR = APP_DIR / "assets"
CUSTOMERS_FILE = DATA_DIR / "customers.json"
DOCUMENTS_FILE = DATA_DIR / "documents.json"
SETTINGS_FILE = DATA_DIR / "settings.json"
LOGO_PATH = ASSETS_DIR / "logo.png"
STAMP_PATH = ASSETS_DIR / "stamp.png"

DATA_DIR.mkdir(exist_ok=True)
ASSETS_DIR.mkdir(exist_ok=True)

COMPANY = {
    "name": "ARTE DI CASA - F.Z.E",
    "address": "Ajman Free Zone C1 Building, Premises Number B.C. 1302226, Ajman Free Zone, Ajman, United Arab Emirates",
    "license": "LICENSE / REGISTRATION NO. 41814",
    "trn": "TRN: 105089220500003",
    "phone": "+39 3333920771 / +971 55 737 7933",
    "email": "operations@artedicasaae.co",
}

VAT_RATE = 0.05

BANKS = {
    "EUR": [
        {"Bank Name":"WIO BANK PJSC","Currency":"EUR","IBAN":"AE300860000009539228335","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE, P.O. Box: 112437"},
    ],
    "USD": [
        {"Bank Name":"WIO BANK PJSC","Currency":"USD","IBAN":"AE310860000009738781276","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE, P.O. Box: 112437"},
    ],
    "AED": [
        {"Bank Name":"WIO BANK PJSC","Currency":"AED","IBAN":"","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE, P.O. Box: 112437"},
    ],
}

PRODUCT_COLS = ["Brand","Product Details","Size","Finish","Qty","Rate Per Piece"]
PACK_COLS = ["Box No","Part","Brand","Product Details","Length","Breadth","Height","CBM","GW","NW"]


def product_row_id():
    return str(uuid.uuid4())


def prepare_product_rows(rows):
    prepared = []
    for row in rows or []:
        new_row = dict(row)
        new_row.setdefault("_row_id", product_row_id())
        for col in PRODUCT_COLS:
            if col not in new_row:
                new_row[col] = 0.0 if col in ["Qty", "Rate Per Piece"] else ""
        prepared.append(new_row)
    if not prepared:
        prepared.append({"_row_id": product_row_id(), "Brand":"", "Product Details":"", "Size":"", "Finish":"", "Qty":1.0, "Rate Per Piece":0.0})
    return prepared


def strip_product_rows(rows):
    cleaned = []
    for row in rows or []:
        item = {col: row.get(col, 0.0 if col in ["Qty", "Rate Per Piece"] else "") for col in PRODUCT_COLS}
        if str(item.get("Brand", "")).strip() or str(item.get("Product Details", "")).strip() or float(item.get("Qty", 0) or 0) or float(item.get("Rate Per Piece", 0) or 0):
            cleaned.append(item)
    return cleaned

def load_json(path, default):
    if not path.exists():
        path.write_text(json.dumps(default, indent=2))
        return default
    try:
        return json.loads(path.read_text())
    except Exception:
        return default

def save_json(path, data):
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False))

def img_b64(path):
    return base64.b64encode(path.read_bytes()).decode() if path.exists() else ""

def money(v, currency):
    symbol = {"EUR":"€","USD":"$","AED":"AED"}.get(currency, currency)
    return f"{symbol} {float(v or 0):,.2f}"

def number_prefix(doc_type):
    return "ADC/PI" if doc_type == "Proforma Invoice" else "ADC/INV"

def next_number(doc_type, docs):
    prefix = number_prefix(doc_type)
    year = datetime.now().year
    nums = []
    for d in docs:
        n = str(d.get("number",""))
        if n.startswith(f"{prefix}/{year}/"):
            try:
                nums.append(int(n.split("/")[-1]))
            except Exception:
                pass
    return f"{prefix}/{year}/{(max(nums) if nums else 0)+1:04d}"

def calculate(products, discount_type, discount_value, shipping_enabled, shipping_cost, vat_mode="VAT 5%"):
    subtotal = 0.0
    for p in products:
        subtotal += float(p.get("Qty",0) or 0) * float(p.get("Rate Per Piece",0) or 0)
    discount = subtotal * float(discount_value or 0) / 100 if discount_type == "Percentage" else float(discount_value or 0)
    shipping = float(shipping_cost or 0) if shipping_enabled else 0.0
    taxable_base = subtotal - discount + shipping
    vat_amount = round(taxable_base * VAT_RATE, 2) if vat_mode == "VAT 5%" else 0.0
    return subtotal, discount, shipping, vat_amount, taxable_base + vat_amount


def is_real_packing_row(row):
    return bool(str(row.get("Brand", "")).strip()) or bool(str(row.get("Product Details", "")).strip())

def clean_packing_rows(rows):
    cleaned = []
    for row in rows or []:
        if not is_real_packing_row(row):
            continue
        cleaned.append(row)
    for idx, row in enumerate(cleaned, 1):
        row["Box No"] = idx
    return cleaned


def packing_row_key(row):
    return f"{row.get('Box No','')}|{row.get('Part','')}|{row.get('Brand','')}|{row.get('Product Details','')}"

def merge_packing_values(old_rows, edited_rows):
    """Preserve manually typed Length/Breadth/Height/GW/NW values during Streamlit reruns."""
    old_by_key = {packing_row_key(r): r for r in old_rows or []}
    merged = []
    for row in edited_rows or []:
        if not is_real_packing_row(row):
            continue
        key = packing_row_key(row)
        old = old_by_key.get(key, {})
        new_row = dict(row)

        # If Streamlit temporarily sends blank/zero during rerun, keep old manual value.
        for field in ["Length", "Breadth", "Height", "GW", "NW"]:
            val = new_row.get(field, "")
            old_val = old.get(field, "")
            if (val == "" or val is None) and old_val not in ["", None]:
                new_row[field] = old_val

        l = float(new_row.get("Length", 0) or 0)
        b = float(new_row.get("Breadth", 0) or 0)
        h = float(new_row.get("Height", 0) or 0)
        new_row["CBM"] = round(l * b * h / 1000000, 3)
        merged.append(new_row)

    return clean_packing_rows(merged)

def packing_summary(rows):
    real_rows = clean_packing_rows(rows)
    return {
        "Total Boxes": len(real_rows),
        "Total CBM": round(sum(float(x.get("CBM", 0) or 0) for x in real_rows), 3),
        "Total GW": round(sum(float(x.get("GW", 0) or 0) for x in real_rows), 2),
        "Total NW": round(sum(float(x.get("NW", 0) or 0) for x in real_rows), 2),
    }

def packing_from_products(products, existing=None):
    existing = existing or []
    out = []

    # Existing packing rows are preserved so split boxes/parts are not lost on edit.
    for i, p in enumerate(products):
        matching = [
            row for row in existing
            if row.get("Product Details", "") == p.get("Product Details", "")
            and row.get("Brand", "") == p.get("Brand", "")
        ]

        if matching:
            for row in matching:
                l = float(row.get("Length", 0) or 0)
                b = float(row.get("Breadth", 0) or 0)
                h = float(row.get("Height", 0) or 0)
                box_qty = float(row.get("Box Qty", 1) or 1)
                row["Box No"] = int(row.get("Box No", len(out) + 1) or len(out) + 1)
                row["Part"] = row.get("Part", "1/1")
                row["Brand"] = p.get("Brand", "")
                row["Product Details"] = p.get("Product Details", "")
                row["CBM"] = round(l * b * h * box_qty / 1000000, 3)
                row["GW"] = float(row.get("GW", 0) or 0)
                row["NW"] = float(row.get("NW", 0) or 0)
                out.append(row)
        else:
            out.append({
                "Box No": len(out) + 1,
                "Part": "1/1",
                "Brand": p.get("Brand", ""),
                "Product Details": p.get("Product Details", ""),
                "Length": 0.0,
                "Breadth": 0.0,
                "Height": 0.0,
                "CBM": 0.0,
                "GW": 0.0,
                "NW": 0.0,
            })

    return clean_packing_rows(out)


def pdf_header_footer(canvas, doc, title=""):
    w, h = A4
    canvas.saveState()

    # Logo
    if LOGO_PATH.exists():
        canvas.drawImage(
            str(LOGO_PATH),
            12 * mm,
            h - 24 * mm,
            width=35 * mm,
            height=16 * mm,
            preserveAspectRatio=True,
            mask="auto"
        )

    # Header text
    canvas.setFillColor(colors.HexColor("#2f80c4"))
    canvas.setFont("Helvetica-Bold", 11)
    canvas.drawRightString(w - 12 * mm, h - 15 * mm, COMPANY["name"])

    canvas.setFillColor(colors.HexColor("#5b9bd5"))
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawRightString(w - 12 * mm, h - 20 * mm, title)

    canvas.setStrokeColor(colors.HexColor("#7db5e8"))
    canvas.line(12 * mm, h - 27 * mm, w - 12 * mm, h - 27 * mm)

    # Stamp bottom right
    if STAMP_PATH.exists():
        try:
            canvas.saveState()
            canvas.setFillAlpha(0.55)
            canvas.drawImage(
                str(STAMP_PATH),
                w - 48 * mm,
                15 * mm,
                width=32 * mm,
                height=32 * mm,
                preserveAspectRatio=True,
                mask="auto"
            )
            canvas.restoreState()
        except Exception:
            canvas.drawImage(
                str(STAMP_PATH),
                w - 48 * mm,
                15 * mm,
                width=32 * mm,
                height=32 * mm,
                preserveAspectRatio=True,
                mask="auto"
            )

    # Footer
    canvas.setFillColor(colors.grey)
    canvas.setFont("Helvetica", 7)
    canvas.drawCentredString(
        w / 2,
        8 * mm,
        f"{COMPANY['name']} | {COMPANY['trn']} | {COMPANY['email']} | Page {canvas.getPageNumber()}"
    )

    canvas.restoreState()




ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

def number_to_words(n):
    n = int(n)
    if n == 0:
        return "Zero"
    if n < 20:
        return ONES[n]
    if n < 100:
        return TENS[n // 10] + ((" " + ONES[n % 10]) if n % 10 else "")
    if n < 1000:
        return ONES[n // 100] + " Hundred" + ((" " + number_to_words(n % 100)) if n % 100 else "")
    if n < 1000000:
        return number_to_words(n // 1000) + " Thousand" + ((" " + number_to_words(n % 1000)) if n % 1000 else "")
    if n < 1000000000:
        return number_to_words(n // 1000000) + " Million" + ((" " + number_to_words(n % 1000000)) if n % 1000000 else "")
    return number_to_words(n // 1000000000) + " Billion" + ((" " + number_to_words(n % 1000000000)) if n % 1000000000 else "")

def amount_in_words(amount, currency):
    amount = round(float(amount or 0), 2)
    whole = int(amount)
    cents = int(round((amount - whole) * 100))
    currency_name = {"EUR": "Euros", "USD": "US Dollars", "AED": "UAE Dirhams"}.get(currency, currency)
    minor_name = {"EUR": "Cents", "USD": "Cents", "AED": "Fils"}.get(currency, "Cents")
    words = f"{number_to_words(whole)} {currency_name}"
    if cents:
        words += f" and {number_to_words(cents)} {minor_name}"
    return words + " Only"


def build_excel(docdata):
    buffer = BytesIO()
    products = docdata.get("products", [])
    packing = packing_from_products(products, docdata.get("packing", []))
    subtotal, disc, shipc, vat_amount, total = calculate(
        products,
        docdata.get("discount_type", "Percentage"),
        docdata.get("discount_value", 0),
        docdata.get("shipping_enabled", False),
        docdata.get("shipping_cost", 0),
        docdata.get("vat_mode", "VAT 5%"),
    )

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        info_rows = [
            ["Document Type", docdata.get("type", "")],
            ["Document No", docdata.get("number", "")],
            ["Date", docdata.get("date", "")],
            ["Currency", docdata.get("currency", "")],
            ["Seller VAT / TRN", COMPANY["trn"].replace("TRN: ", "")],
            ["Customer", docdata.get("bill_to", {}).get("Company Name", "")],
            ["VAT Treatment", docdata.get("vat_mode", "VAT 5%")],
            ["VAT Amount", vat_amount],
            ["Grand Total", total],
        ]
        pd.DataFrame(info_rows, columns=["Field", "Value"]).to_excel(writer, sheet_name="Document Info", index=False)

        invoice_rows = []
        for i, p in enumerate(products, 1):
            qty = float(p.get("Qty", 0) or 0)
            rate = float(p.get("Rate Per Piece", 0) or 0)
            invoice_rows.append({
                "SL": i,
                "Brand": p.get("Brand", ""),
                "Product Details": p.get("Product Details", ""),
                "Size": p.get("Size", ""),
                "Finish": p.get("Finish", ""),
                "Qty": qty,
                "Rate Per Piece": rate,
                "Total Qty": qty,
                "Amount": qty * rate,
            })
        pd.DataFrame(invoice_rows).to_excel(writer, sheet_name="Invoice", index=False)

        total_rows = [
            {"Description": "Subtotal", "Amount": subtotal},
            {"Description": "Discount", "Amount": disc},
            {"Description": "Shipping", "Amount": shipc},
            {"Description": "VAT 5%" if docdata.get("vat_mode", "VAT 5%") == "VAT 5%" else "OUT OF SCOPE OF VAT", "Amount": vat_amount},
            {"Description": "Grand Total", "Amount": total},
        ]
        pd.DataFrame(total_rows).to_excel(writer, sheet_name="Totals", index=False)

        pd.DataFrame(BANKS.get(docdata.get("currency", "EUR"), [])).to_excel(writer, sheet_name="Bank Details", index=False)

        pd.DataFrame([{"Terms & Conditions": docdata.get("terms", "")}]).to_excel(writer, sheet_name="Terms", index=False)

        if docdata.get("type") == "Invoice":
            pd.DataFrame(packing).to_excel(writer, sheet_name="Packing List", index=False)
            summary = [{
                "Total Boxes": len(packing),
                "Total CBM": sum(float(x.get("CBM", 0) or 0) for x in packing),
                "Total GW": sum(float(x.get("GW", 0) or 0) for x in packing),
                "Total NW": sum(float(x.get("NW", 0) or 0) for x in packing),
            }]
            pd.DataFrame(summary).to_excel(writer, sheet_name="Packing Summary", index=False)

        for sheet in writer.sheets.values():
            for col in sheet.columns:
                max_len = 12
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        max_len = max(max_len, len(str(cell.value or "")) + 2)
                    except Exception:
                        pass
                sheet.column_dimensions[col_letter].width = min(max_len, 45)

    buffer.seek(0)
    return buffer.getvalue()



def build_word(docdata):
    buffer = BytesIO()
    products = docdata.get("products", [])
    packing = packing_from_products(products, docdata.get("packing", []))
    subtotal, disc, shipc, vat_amount, total = calculate(
        products,
        docdata.get("discount_type", "Percentage"),
        docdata.get("discount_value", 0),
        docdata.get("shipping_enabled", False),
        docdata.get("shipping_cost", 0),
        docdata.get("vat_mode", "VAT 5%"),
    )

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.8))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(docdata.get("type", "").upper())
    r.bold = True
    r.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"No: {docdata.get('number','')} | Date: {docdata.get('date','')} | Currency: {docdata.get('currency','')}").bold = True

    document.add_paragraph(f"{COMPANY['name']}\n{COMPANY['address']}\n{COMPANY['license']}\n{COMPANY['trn']}\n{COMPANY['phone']}\n{COMPANY['email']}")

    bill = docdata.get("bill_to", {})
    ship = docdata.get("ship_to", {})
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.cell(0,0).text = "BILL TO\n" + "\n".join([
        bill.get("Company Name",""), bill.get("Registration Number",""), bill.get("GST/VAT",""),
        bill.get("Contact Person",""), bill.get("Phone",""), bill.get("Email",""),
        bill.get("Address",""), bill.get("Country","")
    ])
    table.cell(0,1).text = "SHIP TO\n" + ("Same as Bill To" if docdata.get("ship_same") else "\n".join([
        ship.get("Company Name",""), ship.get("Contact Person",""), ship.get("Phone",""),
        ship.get("Email",""), ship.get("Address",""), ship.get("Country","")
    ]))

    document.add_paragraph("")
    prod_table = document.add_table(rows=1, cols=8)
    prod_table.style = "Table Grid"
    headers = ["SL", "Brand", "Product Details", "Size", "Finish", "Qty", "Rate/PC", "Amount"]
    for i, h in enumerate(headers):
        prod_table.rows[0].cells[i].text = h
    for idx, p in enumerate(products, 1):
        qty = float(p.get("Qty", 0) or 0)
        rate = float(p.get("Rate Per Piece", 0) or 0)
        row = prod_table.add_row().cells
        vals = [idx, p.get("Brand",""), p.get("Product Details",""), p.get("Size",""), p.get("Finish",""), qty, money(rate, docdata.get("currency","EUR")), money(qty*rate, docdata.get("currency","EUR"))]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    document.add_paragraph("")
    totals_table = document.add_table(rows=5, cols=2)
    totals_table.style = "Table Grid"
    rows = [
        ("Subtotal", money(subtotal, docdata.get("currency","EUR"))),
        ("Discount", f"- {money(disc, docdata.get('currency','EUR'))}"),
        ("Shipping", money(shipc, docdata.get("currency","EUR"))),
        ("VAT 5%" if docdata.get("vat_mode", "VAT 5%") == "VAT 5%" else "OUT OF SCOPE OF VAT", money(vat_amount, docdata.get("currency","EUR")) if docdata.get("vat_mode", "VAT 5%") == "VAT 5%" else "0.00"),
        ("Grand Total", money(total, docdata.get("currency","EUR"))),
    ]
    for i, (label, val) in enumerate(rows):
        totals_table.cell(i,0).text = label
        totals_table.cell(i,1).text = val

    p = document.add_paragraph()
    r = p.add_run("Amount in Words: ")
    r.bold = True
    p.add_run(amount_in_words(total, docdata.get("currency","EUR")))

    document.add_paragraph("")
    document.add_heading("Bank Details", level=2)
    bank_table = document.add_table(rows=1, cols=5)
    bank_table.style = "Table Grid"
    for i, h in enumerate(["Bank Name", "Currency", "IBAN", "SWIFT/BIC", "Bank Address"]):
        bank_table.rows[0].cells[i].text = h
    for b in BANKS.get(docdata.get("currency","EUR"), []):
        row = bank_table.add_row().cells
        for i, h in enumerate(["Bank Name", "Currency", "IBAN", "SWIFT/BIC", "Bank Address"]):
            row[i].text = b.get(h, "")

    document.add_page_break()
    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
    document.add_heading("Terms & Conditions", level=1)
    document.add_paragraph(docdata.get("terms", ""))

    document.add_paragraph("")
    sig = document.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0,0).text = "Seller Signature"
    sig.cell(0,1).text = "Buyer Signature"
    sig.cell(1,0).text = "HARSH TEJPAL RANA\nOWNER\nARTE DI CASA - F.Z.E"
    sig.cell(1,1).text = "Name:\nCompany:\nDate:"
    if STAMP_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(STAMP_PATH), width=Inches(1.0))

    if docdata.get("type") == "Invoice":
        document.add_page_break()
        if LOGO_PATH.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
        document.add_heading("Packing List", level=1)
        pack_table = document.add_table(rows=1, cols=10)
        pack_table.style = "Table Grid"
        for i, h in enumerate(["SL", "Box No", "Part", "Brand", "Product Details", "L", "B", "H", "CBM", "GW/NW"]):
            pack_table.rows[0].cells[i].text = h
        for i, p in enumerate(packing, 1):
            row = pack_table.add_row().cells
            vals = [i, p.get("Box No", i), p.get("Part",""), p.get("Brand",""), p.get("Product Details",""), p.get("Length",0), p.get("Breadth",0), p.get("Height",0), p.get("CBM",0), f"{p.get('GW',0)} / {p.get('NW',0)}"]
            for j, v in enumerate(vals):
                row[j].text = str(v)
        summary = docdata.get("packing_summary") or packing_summary(packing)
        document.add_paragraph(
            f"Total Boxes: {int(summary.get('Total Boxes', len(packing)))} | "
            f"Total CBM: {float(summary.get('Total CBM', 0)):.3f} | "
            f"Total GW: {float(summary.get('Total GW', 0)):.2f} KG | "
            f"Total NW: {float(summary.get('Total NW', 0)):.2f} KG"
        )
        if STAMP_PATH.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_picture(str(STAMP_PATH), width=Inches(1.0))

    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_pdf(docdata):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=12*mm, rightMargin=12*mm, topMargin=32*mm, bottomMargin=15*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=7.5, leading=9))
    styles.add(ParagraphStyle(name="Tiny", parent=styles["Normal"], fontSize=6.8, leading=8))
    styles.add(ParagraphStyle(name="TitleGold", parent=styles["Title"], fontSize=18, textColor=colors.HexColor("#5b9bd5"), leading=21))
    styles.add(ParagraphStyle(name="Navy", parent=styles["Heading2"], fontSize=11, textColor=colors.HexColor("#2f80c4"), leading=13))
    story = []

    story.append(Paragraph(docdata["type"].upper(), styles["TitleGold"]))
    story.append(Paragraph(f"No: {docdata['number']} &nbsp;&nbsp; Date: {docdata['date']} &nbsp;&nbsp; Currency: {docdata['currency']}", styles["Small"]))
    story.append(Paragraph(f"<b>SELLER VAT / TRN:</b> {COMPANY['trn'].replace('TRN: ', '')}", styles["Navy"]))
    story.append(Paragraph(f"{COMPANY['address']}<br/>{COMPANY['license']}<br/>{COMPANY['trn']}<br/>{COMPANY['phone']}<br/>{COMPANY['email']}", styles["Small"]))
    story.append(Spacer(1,5))

    bill = docdata.get("bill_to",{})
    ship = docdata.get("ship_to",{})
    bill_text = "<br/>".join([f"<b>{bill.get('Company Name','')}</b>", bill.get("Registration Number",""), bill.get("GST/VAT",""), bill.get("Contact Person",""), bill.get("Phone",""), bill.get("Email",""), bill.get("Address",""), bill.get("Country","")])
    ship_text = "Same as Bill To" if docdata.get("ship_same") else "<br/>".join([ship.get("Company Name",""), ship.get("Contact Person",""), ship.get("Phone",""), ship.get("Email",""), ship.get("Address",""), ship.get("Country","")])
    bt = Table([[Paragraph("<b>BILL TO</b><br/>"+bill_text, styles["Small"]), Paragraph("<b>SHIP TO</b><br/>"+ship_text, styles["Small"])]], colWidths=[86*mm,86*mm])
    bt.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#9bc2e6")),("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#f3f9ff")),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(bt); story.append(Spacer(1,6))

    rows = [["SL","Brand","Product Details","Size","Finish","Qty","Rate/PC","Amount"]]
    for i,p in enumerate(docdata["products"], 1):
        qty=float(p.get("Qty",0) or 0); rate=float(p.get("Rate Per Piece",0) or 0)
        rows.append([i, p.get("Brand",""), Paragraph(p.get("Product Details",""), styles["Tiny"]), p.get("Size",""), Paragraph(p.get("Finish",""), styles["Tiny"]), qty, money(rate, docdata["currency"]), money(qty*rate, docdata["currency"])])
    t = Table(rows, repeatRows=1, colWidths=[8*mm,24*mm,43*mm,25*mm,34*mm,10*mm,23*mm,24*mm])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#2f80c4")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("GRID",(0,0),(-1,-1),0.25,colors.lightgrey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(t); story.append(Spacer(1,6))

    subtotal, disc, shipc, vat_amount, total = calculate(docdata["products"], docdata["discount_type"], docdata["discount_value"], docdata["shipping_enabled"], docdata["shipping_cost"], docdata.get("vat_mode", "VAT 5%"))
    totals = [["Subtotal", money(subtotal, docdata["currency"])],["Discount", f"- {money(disc, docdata['currency'])}"],["Shipping", money(shipc, docdata["currency"])],["VAT 5%" if docdata.get("vat_mode", "VAT 5%") == "VAT 5%" else "OUT OF SCOPE OF VAT", money(vat_amount, docdata["currency"]) if docdata.get("vat_mode", "VAT 5%") == "VAT 5%" else "0.00"],["Grand Total", money(total, docdata["currency"])]]
    tt = Table(totals, colWidths=[40*mm,38*mm], hAlign="RIGHT")
    tt.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.25,colors.grey),("BACKGROUND",(0,4),(-1,4),colors.HexColor("#2f80c4")),("TEXTCOLOR",(0,4),(-1,4),colors.white),("FONTNAME",(0,4),(-1,4),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8)]))
    story.append(tt); story.append(Spacer(1,5))
    story.append(Paragraph(f"<b>Amount in Words:</b> {amount_in_words(total, docdata['currency'])}", styles["Navy"]))
    story.append(Spacer(1,7))

    story.append(Paragraph("BANK DETAILS", styles["Navy"]))
    bank_rows = [["Bank Name","Currency","IBAN","SWIFT/BIC","Bank Address"]]
    for b in BANKS[docdata["currency"]]:
        bank_rows.append([b["Bank Name"], b["Currency"], b["IBAN"], b["SWIFT/BIC"], Paragraph(b["Bank Address"], styles["Tiny"])])
    bank_table = Table(bank_rows, repeatRows=1, colWidths=[38*mm,16*mm,45*mm,25*mm,55*mm])
    bank_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#5b9bd5")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.25,colors.grey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(bank_table)

    story.append(PageBreak())
    story.append(Paragraph("TERMS & CONDITIONS", styles["TitleGold"]))
    if LOGO_PATH.exists(): story.append(Spacer(1,2))
    for line in docdata.get("terms","").split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), styles["Small"]))
            story.append(Spacer(1,1.5))
    story.append(Spacer(1,12))
    sig = [["Seller Signature", "Buyer Signature"]]
    if STAMP_PATH.exists():
        sig.append([Image(str(STAMP_PATH), width=28*mm, height=20*mm), ""])
    sig.append(["HARSH TEJPAL RANA\nOWNER\nARTE DI CASA - F.Z.E", "Name:\nCompany:\nDate:"])
    sigt = Table(sig, colWidths=[85*mm,85*mm])
    sigt.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.25,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#f3f9ff")),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),8)]))
    story.append(sigt)

    if docdata["type"] == "Invoice":
        story.append(PageBreak())
        story.append(Paragraph("PACKING LIST", styles["TitleGold"]))
        pack = packing_from_products(docdata["products"], docdata.get("packing",[]))
        prow = [["SL","Box No","Part","Brand","Product Details","Length","Breadth","Height","CBM","GW","NW"]]
        for i,p in enumerate(pack, 1):
            prow.append([i,p.get("Box No", i),p.get("Part","1/1"),p["Brand"],Paragraph(p["Product Details"],styles["Tiny"]),p["Length"],p["Breadth"],p["Height"],p["CBM"],p["GW"],p["NW"]])
        pt = Table(prow, repeatRows=1, colWidths=[7*mm,13*mm,13*mm,24*mm,45*mm,13*mm,13*mm,13*mm,16*mm,16*mm,16*mm])
        pt.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#2f80c4")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.25,colors.lightgrey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
        story.append(pt); story.append(Spacer(1,8))
        summary = docdata.get("packing_summary") or packing_summary(pack)
        story.append(Paragraph(f"<b>Total Boxes:</b> {int(summary.get('Total Boxes', len(pack)))} &nbsp;&nbsp; <b>Total CBM:</b> {float(summary.get('Total CBM', 0)):.3f} &nbsp;&nbsp; <b>Total GW:</b> {float(summary.get('Total GW', 0)):.2f} KG &nbsp;&nbsp; <b>Total NW:</b> {float(summary.get('Total NW', 0)):.2f} KG", styles["Navy"]))
        story.append(Spacer(1,15))
        story.append(Paragraph("Digital Signature / Authorized Signatory", styles["Navy"]))
        if STAMP_PATH.exists():
            story.append(Image(str(STAMP_PATH), width=30*mm, height=22*mm))
        story.append(Paragraph("HARSH TEJPAL RANA<br/>MANAGER<br/>ARTE DI CASA - F.Z.E", styles["Small"]))

    pdf.build(story, onFirstPage=lambda c,d: pdf_header_footer(c,d,docdata["type"]), onLaterPages=lambda c,d: pdf_header_footer(c,d,docdata["type"]))
    return buffer.getvalue()


def parse_amount_value(value):
    """Return a float from values such as € 1,234.50, AED 1,234 or 1,234."""
    txt = str(value or "").replace("€", "").replace("$", "").replace("AED", "").replace(",", "").strip()
    cleaned = "".join(ch for ch in txt if ch.isdigit() or ch in ".-")
    try:
        return float(cleaned) if cleaned not in ["", ".", "-"] else 0.0
    except Exception:
        return 0.0


def parse_gw_nw(value):
    """Return gross/net weights from combined values such as '12 / 10' or 'GW 12 NW 10'."""
    txt = str(value or "").replace("KG", "").replace("kg", "").strip()
    if "/" in txt:
        left, right = txt.split("/", 1)
        return parse_amount_value(left), parse_amount_value(right)
    nums = []
    current = ""
    for ch in txt:
        if ch.isdigit() or ch in ".-":
            current += ch
        elif current:
            try:
                nums.append(float(current))
            except Exception:
                pass
            current = ""
    if current:
        try:
            nums.append(float(current))
        except Exception:
            pass
    if len(nums) >= 2:
        return nums[0], nums[1]
    if len(nums) == 1:
        return nums[0], 0.0
    return 0.0, 0.0


def parse_word_upload(uploaded_file, docs):
    """Import an extracted DOCX/proforma/invoice into the app's editable document structure."""
    document = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    full_text = "\n".join(paragraphs)
    upper_text = full_text.upper()

    doc_type = "Invoice" if "INVOICE" in upper_text and "PROFORMA" not in upper_text else "Proforma Invoice"
    currency = "EUR"
    for cur in ["EUR", "USD", "AED"]:
        if cur in upper_text or (cur == "EUR" and "€" in full_text) or (cur == "USD" and "$" in full_text):
            currency = cur
            break

    imported = {
        "id": str(uuid.uuid4()),
        "type": doc_type,
        "number": next_number(doc_type, docs),
        "date": str(date.today()),
        "currency": currency,
        "bill_to": {"Company Name":"", "Registration Number":"", "GST/VAT":"", "Contact Person":"", "Phone":"", "Email":"", "Country":"", "Address":""},
        "ship_same": True,
        "ship_to": {},
        "products": [],
        "discount_type": "Flat Amount",
        "discount_value": 0.0,
        "shipping_enabled": False,
        "shipping_cost": 0.0,
        "vat_mode": "VAT 5%",
        "vat_rate": VAT_RATE,
        "vat_amount": 0.0,
        "terms": settings.get("terms", "") if 'settings' in globals() else "",
        "packing": [],
        "packing_summary": {"Total Boxes": 0, "Total CBM": 0.0, "Total GW": 0.0, "Total NW": 0.0},
        "total": 0.0,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }

    # Document number, date and currency from text if the exported file contains them.
    for line in paragraphs:
        low = line.lower()
        if "no:" in low or "invoice no" in low or "document no" in low:
            parts = line.replace("|", " ").split()
            for i, token in enumerate(parts):
                if token.lower().rstrip(":") in ["no", "number", "no:"] and i + 1 < len(parts):
                    candidate = parts[i + 1].strip()
                    if "/" in candidate or "-" in candidate:
                        imported["number"] = candidate
                        break
        if "currency" in low:
            for cur in ["EUR", "USD", "AED"]:
                if cur in line.upper():
                    imported["currency"] = cur
        if "out of scope" in low:
            imported["vat_mode"] = "OUT OF SCOPE OF VAT"

    # Read tables. The app export uses clear product and customer tables; this also handles many extracted Word tables.
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [h.strip().lower() for h in rows[0]]
        header_join = " | ".join(header)

        # Product rows: SL, Brand, Product Details, Size, Finish, Qty, Rate/PC, Amount
        if ("brand" in header_join and ("product" in header_join or "description" in header_join) and ("qty" in header_join or "quantity" in header_join)):
            colmap = {name: idx for idx, name in enumerate(header)}
            def find_col(*names):
                for name in names:
                    for idx, h in enumerate(header):
                        if name in h:
                            return idx
                return None
            brand_i = find_col("brand")
            prod_i = find_col("product details", "description", "product")
            size_i = find_col("size")
            finish_i = find_col("finish")
            qty_i = find_col("qty", "quantity")
            rate_i = find_col("rate", "rate/pc", "rate per piece")
            amount_i = find_col("amount")
            for row in rows[1:]:
                def cell(i): return row[i].strip() if i is not None and i < len(row) else ""
                details = cell(prod_i)
                brand = cell(brand_i)
                if not brand and not details:
                    continue
                qty = parse_amount_value(cell(qty_i)) or 1.0
                rate = parse_amount_value(cell(rate_i))
                amount = parse_amount_value(cell(amount_i))
                if not rate and amount and qty:
                    rate = amount / qty
                imported["products"].append({
                    "Brand": brand,
                    "Product Details": details,
                    "Size": cell(size_i),
                    "Finish": cell(finish_i),
                    "Qty": qty,
                    "Rate Per Piece": rate,
                })

        # Packing list rows: SL, Box No, Part, Brand, Product Details, L/B/H, CBM, GW/NW
        if (("box" in header_join or "box no" in header_join) and "cbm" in header_join) or ("gw" in header_join and "nw" in header_join and "product" in header_join):
            def find_col(*names):
                for name in names:
                    for idx, h in enumerate(header):
                        if name == h or name in h:
                            return idx
                return None

            box_i = find_col("box no", "box")
            part_i = find_col("part")
            brand_i = find_col("brand")
            prod_i = find_col("product details", "description", "product")
            length_i = find_col("length", "l")
            breadth_i = find_col("breadth", "b")
            height_i = find_col("height", "h")
            cbm_i = find_col("cbm")
            gw_i = find_col("gw")
            nw_i = find_col("nw")
            gwnw_i = find_col("gw/nw", "gw / nw")

            imported_packing = []
            for row in rows[1:]:
                def cell(i): return row[i].strip() if i is not None and i < len(row) else ""
                brand = cell(brand_i)
                details = cell(prod_i)
                if not brand and not details:
                    continue
                gw = parse_amount_value(cell(gw_i)) if gw_i is not None and gw_i != gwnw_i else 0.0
                nw = parse_amount_value(cell(nw_i)) if nw_i is not None and nw_i != gwnw_i else 0.0
                if (not gw and not nw) and gwnw_i is not None:
                    gw, nw = parse_gw_nw(cell(gwnw_i))
                length = parse_amount_value(cell(length_i))
                breadth = parse_amount_value(cell(breadth_i))
                height = parse_amount_value(cell(height_i))
                cbm = parse_amount_value(cell(cbm_i))
                if not cbm:
                    cbm = round(length * breadth * height / 1000000, 3)
                imported_packing.append({
                    "Box No": int(parse_amount_value(cell(box_i)) or len(imported_packing) + 1),
                    "Part": cell(part_i) or "1/1",
                    "Brand": brand,
                    "Product Details": details,
                    "Length": length,
                    "Breadth": breadth,
                    "Height": height,
                    "CBM": cbm,
                    "GW": gw,
                    "NW": nw,
                })

            if imported_packing:
                imported["type"] = "Invoice"
                imported["packing"] = clean_packing_rows(imported_packing)
                imported["packing_summary"] = packing_summary(imported["packing"])

        # Bill/ship table from this app's Word export.
        if len(rows) == 1 and len(rows[0]) >= 1 and "bill to" in rows[0][0].lower():
            bill_lines = [x.strip() for x in rows[0][0].splitlines() if x.strip() and x.strip().lower() != "bill to"]
            if bill_lines:
                imported["bill_to"]["Company Name"] = bill_lines[0]
                if len(bill_lines) > 1: imported["bill_to"]["Registration Number"] = bill_lines[1]
                if len(bill_lines) > 2: imported["bill_to"]["GST/VAT"] = bill_lines[2]
                if len(bill_lines) > 3: imported["bill_to"]["Contact Person"] = bill_lines[3]
                if len(bill_lines) > 4: imported["bill_to"]["Phone"] = bill_lines[4]
                if len(bill_lines) > 5: imported["bill_to"]["Email"] = bill_lines[5]
                if len(bill_lines) > 6: imported["bill_to"]["Address"] = bill_lines[6]
                if len(bill_lines) > 7: imported["bill_to"]["Country"] = bill_lines[7]
            if len(rows[0]) > 1 and "same as bill" not in rows[0][1].lower():
                imported["ship_same"] = False
                ship_lines = [x.strip() for x in rows[0][1].splitlines() if x.strip() and x.strip().lower() != "ship to"]
                imported["ship_to"] = {
                    "Company Name": ship_lines[0] if len(ship_lines) > 0 else "",
                    "Contact Person": ship_lines[1] if len(ship_lines) > 1 else "",
                    "Phone": ship_lines[2] if len(ship_lines) > 2 else "",
                    "Email": ship_lines[3] if len(ship_lines) > 3 else "",
                    "Address": ship_lines[4] if len(ship_lines) > 4 else "",
                    "Country": ship_lines[5] if len(ship_lines) > 5 else "",
                }

    # Terms from extracted Word, when available.
    if "TERMS & CONDITIONS" in upper_text:
        try:
            imported["terms"] = full_text.split("TERMS & CONDITIONS", 1)[1].strip() or imported["terms"]
        except Exception:
            pass

    if not imported["products"]:
        imported["products"] = [{"Brand":"", "Product Details":"", "Size":"", "Finish":"", "Qty":1.0, "Rate Per Piece":0.0}]

    subtotal, disc, shipc, vat_amount, total = calculate(
        imported["products"], imported["discount_type"], imported["discount_value"], imported["shipping_enabled"], imported["shipping_cost"], imported["vat_mode"]
    )
    imported["vat_amount"] = vat_amount
    imported["total"] = total
    return imported


settings = load_json(SETTINGS_FILE, {"terms":"","password":"1985"})
st.set_page_config(page_title="Arte Di Casa Billing", layout="wide")

# optional password, on by default
if "auth" not in st.session_state: st.session_state.auth = False
if not st.session_state.auth:
    st.markdown("<div style='background:#2f80c4;padding:38px;border-radius:22px;text-align:center;color:white;margin-top:80px;'><h1 style='color:#7db5e8;font-family:Georgia;'>ARTE DI CASA</h1><p style='letter-spacing:4px;'>BILLING SYSTEM LOGIN</p></div>", unsafe_allow_html=True)
    pw = st.text_input("Password", type="password")
    if st.button("Login"):
        if pw == str(settings.get("password","1985")):
            st.session_state.auth = True
            st.rerun()
        else:
            st.error("Incorrect password")
    st.stop()

st.markdown("""
<style>
.stApp { background:#f7f3ec; }
.block-container { max-width:1600px; padding-top:1rem; }
div[data-testid="stButton"] button { font-size: 11px !important; padding: 0.25rem 0.35rem !important; min-height: 30px !important; }

.cap { background:#2f80c4; color:white; padding:22px; border-radius:22px; margin-bottom:18px; box-shadow:0 12px 30px #0002; }
.gold { color:#7db5e8; font-family:Georgia,serif; }
.card { background:white; border:1px solid #d7e9fb; border-radius:18px; padding:16px; margin-bottom:16px; box-shadow:0 8px 24px #2f80c410; }
</style>
""", unsafe_allow_html=True)

logo = img_b64(LOGO_PATH)
logo_html = f"<img src='data:image/png;base64,{logo}' style='height:78px;object-fit:contain;margin-right:18px;'>" if logo else ""
st.markdown(f"<div class='cap' style='display:flex;align-items:center;'>{logo_html}<div><h1 class='gold' style='margin:0;font-size:34px;'>ARTE DI CASA</h1><div style='letter-spacing:5px;font-size:11px;'>BILLING · PROFORMA · INVOICE · PACKING LIST</div></div></div>", unsafe_allow_html=True)

customers = load_json(CUSTOMERS_FILE, [])
documents = load_json(DOCUMENTS_FILE, [])

if "editing_id" not in st.session_state: st.session_state.editing_id = None



PAGES = ["Create / Edit", "Saved Documents", "Customers", "Settings"]

if "page" not in st.session_state:
    st.session_state.page = "Create / Edit"

# Hidden navigation intent used by Edit / Convert / Save actions.
if "force_page" in st.session_state:
    st.session_state.page = st.session_state.force_page
    del st.session_state.force_page

# Compact top navigation - no wide sidebar.
nav_cols = st.columns([1, 1, 1, 1, 6])
if nav_cols[0].button("Create / Edit", use_container_width=True):
    st.session_state.page = "Create / Edit"
    st.rerun()
if nav_cols[1].button("Saved", use_container_width=True):
    st.session_state.page = "Saved Documents"
    st.rerun()
if nav_cols[2].button("Customers", use_container_width=True):
    st.session_state.page = "Customers"
    st.rerun()
if nav_cols[3].button("Settings", use_container_width=True):
    st.session_state.page = "Settings"
    st.rerun()

st.caption(f"Current page: {st.session_state.page}")

if st.session_state.page == "Create / Edit":
    editing = next((d for d in documents if d.get("id") == st.session_state.editing_id), None) if st.session_state.editing_id else None
    imported_word_doc = st.session_state.get("imported_word_docdata") if not st.session_state.editing_id else None
    if imported_word_doc:
        editing = imported_word_doc
        st.success("Imported Word document loaded. You can edit it, save it, convert it to invoice, and add packing details.")
    elif editing:
        st.success(f"Editing existing document: {editing.get('number','')}. Saving will update the SAME document.")
    else:
        st.info("Creating a new document.")

    c1,c2,c3,c4 = st.columns(4)
    doc_type = c1.selectbox("Document Type", ["Proforma Invoice","Invoice"], index=1 if editing and editing.get("type")=="Invoice" else 0)
    currency = c2.selectbox("Currency", ["EUR","USD","AED"], index=["EUR","USD","AED"].index(editing.get("currency","EUR")) if editing else 0)
    doc_date = c3.date_input("Date", value=datetime.strptime(editing.get("date"), "%Y-%m-%d").date() if editing else date.today())
    doc_number = c4.text_input("Document No.", value=editing.get("number") if editing else next_number(doc_type, documents))

    if st.button("Start New Blank Document"):
        st.session_state.editing_id = None
        st.session_state.pop("imported_word_docdata", None)
        st.session_state.pop("product_rows_new_document", None)
        st.session_state.pop("packing_rows_new_document", None)
        st.session_state.pop("active_product_doc_key", None)
        st.session_state.pop("active_packing_doc_key", None)
        st.session_state.page = "Create / Edit"
        st.rerun()

    st.markdown("<div class='card'><h3 class='gold'>Import Extracted Word File</h3>", unsafe_allow_html=True)
    uploaded_word = st.file_uploader("Upload extracted Word file (.docx) to bring it into the editor", type=["docx"], key="word_import_upload")
    import_cols = st.columns([1, 5])
    if import_cols[0].button("Import Word", key="import_word_button"):
        if uploaded_word is None:
            st.warning("Upload a .docx file first.")
        else:
            try:
                imported_doc = parse_word_upload(uploaded_word, documents)
                st.session_state.imported_word_docdata = imported_doc
                st.session_state.editing_id = None
                st.session_state.pop("active_product_doc_key", None)
                st.session_state.pop("active_packing_doc_key", None)
                st.success("Word file imported. Review and edit the fields below, then save or convert to invoice.")
                st.rerun()
            except Exception as e:
                st.error(f"Could not import this Word file: {e}")
    import_cols[1].caption("Best results: upload a Word file exported from this app or an extracted invoice/proforma with a product table containing Brand/Product/Qty/Rate columns.")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Customer Details</h3>", unsafe_allow_html=True)
    names = ["-- New Customer --"] + [c.get("Company Name","") for c in customers]
    selected = st.selectbox("Choose saved customer", names)
    selected_customer = next((c for c in customers if c.get("Company Name")==selected), {}) if selected != "-- New Customer --" else {}
    bill_existing = editing.get("bill_to", {}) if editing else selected_customer
    a,b = st.columns(2)
    with a:
        bill_company = st.text_input("Company Name", value=bill_existing.get("Company Name",""))
        bill_reg = st.text_input("Company Registration Number", value=bill_existing.get("Registration Number",""))
        bill_vat = st.text_input("GST / VAT", value=bill_existing.get("GST/VAT",""))
        bill_contact = st.text_input("Contact Person", value=bill_existing.get("Contact Person",""))
    with b:
        bill_phone = st.text_input("Phone", value=bill_existing.get("Phone",""))
        bill_email = st.text_input("Email", value=bill_existing.get("Email",""))
        bill_country = st.text_input("Country", value=bill_existing.get("Country",""))
        bill_address = st.text_area("Billing Address", value=bill_existing.get("Address",""), height=92)
    ship_same = st.checkbox("Ship To same as Bill To", value=editing.get("ship_same", True) if editing else True)
    ship_to = {}
    if not ship_same:
        s1,s2 = st.columns(2)
        ship_existing = editing.get("ship_to", {}) if editing else selected_customer.get("ship_to", {})
        with s1:
            ship_to["Company Name"] = st.text_input("Shipping Company", value=ship_existing.get("Company Name",""))
            ship_to["Contact Person"] = st.text_input("Shipping Contact", value=ship_existing.get("Contact Person",""))
            ship_to["Phone"] = st.text_input("Shipping Phone", value=ship_existing.get("Phone",""))
        with s2:
            ship_to["Email"] = st.text_input("Shipping Email", value=ship_existing.get("Email",""))
            ship_to["Country"] = st.text_input("Shipping Country", value=ship_existing.get("Country",""))
            ship_to["Address"] = st.text_area("Shipping Address", value=ship_existing.get("Address",""), height=92)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Products</h3>", unsafe_allow_html=True)
    current_product_doc_key = editing.get("id") if editing else "new_document"
    product_state_key = f"product_rows_{current_product_doc_key}"

    if st.session_state.get("active_product_doc_key") != current_product_doc_key:
        init_products = editing.get("products") if editing else [{"Brand":"","Product Details":"","Size":"","Finish":"","Qty":1,"Rate Per Piece":0.0}]
        st.session_state[product_state_key] = prepare_product_rows(init_products)
        st.session_state["active_product_doc_key"] = current_product_doc_key

    if product_state_key not in st.session_state:
        st.session_state[product_state_key] = prepare_product_rows([])

    st.caption("Delete any product row with X. Serial numbers in invoice/PDF/packing list will automatically move up.")
    prod_headers = st.columns([0.6, 1.4, 3.2, 1.5, 1.5, 0.9, 1.2, 0.7])
    for col, label in zip(prod_headers, ["SL", "Brand", "Product Details", "Size", "Finish", "Qty", "Rate/PC", "Del"]):
        col.markdown(f"**{label}**")

    product_rows = []
    delete_product_index = None
    for idx, row in enumerate(prepare_product_rows(st.session_state[product_state_key])):
        rid = row.get("_row_id", product_row_id())
        c = st.columns([0.6, 1.4, 3.2, 1.5, 1.5, 0.9, 1.2, 0.7])
        c[0].write(idx + 1)
        brand = c[1].text_input("Brand", value=str(row.get("Brand", "")), key=f"prod_brand_{rid}", label_visibility="collapsed")
        details = c[2].text_input("Product Details", value=str(row.get("Product Details", "")), key=f"prod_details_{rid}", label_visibility="collapsed")
        size = c[3].text_input("Size", value=str(row.get("Size", "")), key=f"prod_size_{rid}", label_visibility="collapsed")
        finish = c[4].text_input("Finish", value=str(row.get("Finish", "")), key=f"prod_finish_{rid}", label_visibility="collapsed")
        qty = c[5].number_input("Qty", min_value=0.0, value=float(row.get("Qty", 0) or 0), step=1.0, key=f"prod_qty_{rid}", label_visibility="collapsed")
        rate = c[6].number_input("Rate", min_value=0.0, value=float(row.get("Rate Per Piece", 0) or 0), step=1.0, key=f"prod_rate_{rid}", label_visibility="collapsed")
        if c[7].button("X", key=f"prod_delete_{rid}"):
            delete_product_index = idx
        product_rows.append({"_row_id": rid, "Brand": brand, "Product Details": details, "Size": size, "Finish": finish, "Qty": float(qty or 0), "Rate Per Piece": float(rate or 0)})

    if delete_product_index is not None:
        product_rows.pop(delete_product_index)
        st.session_state[product_state_key] = prepare_product_rows(product_rows)
        st.rerun()

    add_cols = st.columns([1, 5])
    if add_cols[0].button("Add Product Row", key=f"add_product_{current_product_doc_key}"):
        product_rows.append({"_row_id": product_row_id(), "Brand":"", "Product Details":"", "Size":"", "Finish":"", "Qty":1.0, "Rate Per Piece":0.0})
        st.session_state[product_state_key] = prepare_product_rows(product_rows)
        st.rerun()

    st.session_state[product_state_key] = prepare_product_rows(product_rows)
    products = strip_product_rows(st.session_state[product_state_key])
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Discount / Shipping / Totals</h3>", unsafe_allow_html=True)
    d1,d2,d3,d4 = st.columns(4)
    discount_type = d1.selectbox("Discount Type", ["Percentage","Flat Amount"], index=["Percentage","Flat Amount"].index(editing.get("discount_type","Percentage")) if editing else 0)
    discount_value = d2.number_input("Discount Value", min_value=0.0, value=float(editing.get("discount_value",0) if editing else 0))
    shipping_enabled = d3.checkbox("Add Shipping Charges?", value=bool(editing.get("shipping_enabled", False)) if editing else False)
    shipping_cost = d4.number_input("Shipping Cost", min_value=0.0, value=float(editing.get("shipping_cost",0) if editing else 0), disabled=not shipping_enabled, help="Tick Add Shipping Charges first, then enter shipping amount.")
    vat_mode = st.selectbox("VAT Treatment", ["VAT 5%", "OUT OF SCOPE OF VAT"], index=["VAT 5%", "OUT OF SCOPE OF VAT"].index(editing.get("vat_mode", "VAT 5%")) if editing and editing.get("vat_mode") in ["VAT 5%", "OUT OF SCOPE OF VAT"] else 0, help="Use VAT 5% for UAE standard-rated supplies. Use OUT OF SCOPE OF VAT only when your accountant confirms it applies.")
    subtotal, discount_amount, shipping_amount, vat_amount, grand = calculate(products, discount_type, discount_value, shipping_enabled, shipping_cost, vat_mode)
    m1,m2,m3,m4,m5 = st.columns(5)
    m1.metric("Subtotal", money(subtotal, currency))
    m2.metric("Discount", money(discount_amount, currency))
    m3.metric("Shipping", money(shipping_amount, currency))
    m4.metric("VAT", money(vat_amount, currency) if vat_mode == "VAT 5%" else "OUT OF SCOPE")
    m5.metric("Grand Total", money(grand, currency))
    st.info(f"Seller VAT / TRN: {COMPANY['trn'].replace('TRN: ', '')}")
    st.info(f"Amount in Words: {amount_in_words(grand, currency)}")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Bank Details</h3>", unsafe_allow_html=True)
    st.dataframe(pd.DataFrame(BANKS[currency]), use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Terms & Conditions</h3>", unsafe_allow_html=True)
    terms = st.text_area("Editable default terms for this document", value=editing.get("terms", settings.get("terms","")) if editing else settings.get("terms",""), height=330)
    st.markdown("</div>", unsafe_allow_html=True)

    packing = []
    packing_summary_values = {"Total Boxes": 0, "Total CBM": 0.0, "Total GW": 0.0, "Total NW": 0.0}
    packing_manual_override = False

    if doc_type == "Invoice":
        st.markdown("<div class='card'><h3 class='gold'>Packing List — Mandatory</h3>", unsafe_allow_html=True)

        current_doc_key = editing.get("id") if editing else "new_document"
        packing_state_key = f"packing_rows_{current_doc_key}"

        # Load saved packing rows only once per document.
        if st.session_state.get("active_packing_doc_key") != current_doc_key:
            saved_packing = editing.get("packing", []) if editing else []
            st.session_state[packing_state_key] = clean_packing_rows(packing_from_products(products, saved_packing))
            st.session_state["active_packing_doc_key"] = current_doc_key

        if packing_state_key not in st.session_state:
            saved_packing = editing.get("packing", []) if editing else []
            st.session_state[packing_state_key] = clean_packing_rows(packing_from_products(products, saved_packing))

        # Keep packing list aligned with current products. If a product row is deleted,
        # its packing rows are removed and all box numbers move up automatically.
        valid_pairs = {
            (p.get("Brand", ""), p.get("Product Details", ""))
            for p in products
            if str(p.get("Brand", "")).strip() or str(p.get("Product Details", "")).strip()
        }
        rows = [r for r in clean_packing_rows(st.session_state[packing_state_key]) if (r.get("Brand", ""), r.get("Product Details", "")) in valid_pairs]
        existing_pairs = {(r.get("Brand", ""), r.get("Product Details", "")) for r in rows}
        for p in products:
            if not str(p.get("Brand", "")).strip() and not str(p.get("Product Details", "")).strip():
                continue
            pair = (p.get("Brand", ""), p.get("Product Details", ""))
            if pair not in existing_pairs:
                rows.append({
                    "Box No": len(rows) + 1,
                    "Part": "1/1",
                    "Brand": p.get("Brand", ""),
                    "Product Details": p.get("Product Details", ""),
                    "Length": 0.0,
                    "Breadth": 0.0,
                    "Height": 0.0,
                    "CBM": 0.0,
                    "GW": 0.0,
                    "NW": 0.0,
                })
                existing_pairs.add(pair)
        rows = clean_packing_rows(rows)
        st.session_state[packing_state_key] = rows

        st.caption("Stable entry mode: values will not reset while typing. Use Add/Split Box for multiple boxes/parts.")

        product_labels = [
            f"{i+1}. {p.get('Brand','')} - {p.get('Product Details','')}"
            for i, p in enumerate(products)
            if str(p.get("Brand", "")).strip() or str(p.get("Product Details", "")).strip()
        ]

        split_cols = st.columns([3, 1, 1])
        with split_cols[0]:
            split_choice = st.selectbox("Select item to add another box/part", product_labels if product_labels else ["No products"], key=f"split_choice_{current_doc_key}")
        with split_cols[1]:
            part_label = st.text_input("Part label", value="2/2", key=f"part_label_{current_doc_key}")
        with split_cols[2]:
            st.write("")
            st.write("")
            if st.button("Add/Split Box", key=f"add_split_{current_doc_key}"):
                if product_labels:
                    original_index = int(split_choice.split(".")[0]) - 1
                    p = products[original_index]
                    current_rows = clean_packing_rows(st.session_state[packing_state_key])
                    new_split_row = {
                        "Box No": 0,
                        "Part": part_label or "Part",
                        "Brand": p.get("Brand", ""),
                        "Product Details": p.get("Product Details", ""),
                        "Length": 0.0,
                        "Breadth": 0.0,
                        "Height": 0.0,
                        "CBM": 0.0,
                        "GW": 0.0,
                        "NW": 0.0,
                    }
                    insert_at = len(current_rows)
                    for r_idx, existing_row in enumerate(current_rows):
                        if existing_row.get("Brand", "") == p.get("Brand", "") and existing_row.get("Product Details", "") == p.get("Product Details", ""):
                            insert_at = r_idx + 1
                    current_rows.insert(insert_at, new_split_row)
                    st.session_state[packing_state_key] = clean_packing_rows(current_rows)
                    st.rerun()

        st.markdown("#### Packing Rows")
        header_cols = st.columns([0.7, 1.0, 1.0, 1.8, 3.0, 1.0, 1.0, 1.0, 1.0, 1.0, 0.8])
        headers = ["Box", "Part", "Brand", "Product", "Details", "L", "B", "H", "GW", "NW", "Del"]
        for col, h in zip(header_cols, headers):
            col.markdown(f"**{h}**")

        updated_rows = []
        delete_index = None

        for idx, row in enumerate(clean_packing_rows(st.session_state[packing_state_key])):
            row_key = f"{current_doc_key}_{idx}_{row.get('Brand','')}_{row.get('Product Details','')}_{row.get('Part','')}"
            c = st.columns([0.7, 1.0, 1.0, 1.8, 3.0, 1.0, 1.0, 1.0, 1.0, 1.0, 0.8])

            box_no = idx + 1
            c[0].write(box_no)
            part = c[1].text_input("Part", value=str(row.get("Part", "1/1")), key=f"pl_part_{row_key}", label_visibility="collapsed")
            brand = c[2].text_input("Brand", value=str(row.get("Brand", "")), key=f"pl_brand_{row_key}", label_visibility="collapsed")
            product = c[3].text_input("Product", value=str(row.get("Product Details", "")), key=f"pl_product_{row_key}", label_visibility="collapsed")
            details = c[4].caption(str(row.get("Product Details", "")))

            length = c[5].number_input("L", min_value=0.0, value=float(row.get("Length", 0) or 0), step=1.0, key=f"pl_l_{row_key}", label_visibility="collapsed")
            breadth = c[6].number_input("B", min_value=0.0, value=float(row.get("Breadth", 0) or 0), step=1.0, key=f"pl_b_{row_key}", label_visibility="collapsed")
            height = c[7].number_input("H", min_value=0.0, value=float(row.get("Height", 0) or 0), step=1.0, key=f"pl_h_{row_key}", label_visibility="collapsed")
            gw = c[8].number_input("GW", min_value=0.0, value=float(row.get("GW", 0) or 0), step=0.1, key=f"pl_gw_{row_key}", label_visibility="collapsed")
            nw = c[9].number_input("NW", min_value=0.0, value=float(row.get("NW", 0) or 0), step=0.1, key=f"pl_nw_{row_key}", label_visibility="collapsed")

            if c[10].button("X", key=f"pl_del_{row_key}"):
                delete_index = idx

            cbm_value = round(float(length or 0) * float(breadth or 0) * float(height or 0) / 1000000, 3)

            updated_rows.append({
                "Box No": box_no,
                "Part": part,
                "Brand": brand,
                "Product Details": product,
                "Length": float(length or 0),
                "Breadth": float(breadth or 0),
                "Height": float(height or 0),
                "CBM": cbm_value,
                "GW": float(gw or 0),
                "NW": float(nw or 0),
            })

        if delete_index is not None:
            updated_rows.pop(delete_index)
            st.session_state[packing_state_key] = clean_packing_rows(updated_rows)
            st.rerun()

        packing = clean_packing_rows(updated_rows)
        st.session_state[packing_state_key] = packing

        auto_summary = packing_summary(packing)

        st.markdown("**Packing Summary**")
        packing_manual_override = st.checkbox("Manual override summary totals", value=False, key=f"manual_override_{current_doc_key}")
        sc1, sc2, sc3, sc4 = st.columns(4)

        if packing_manual_override:
            total_boxes = sc1.number_input("Total Boxes", min_value=0, value=int(auto_summary["Total Boxes"]), key=f"sum_boxes_{current_doc_key}")
            total_cbm = sc2.number_input("Total CBM", min_value=0.0, value=float(auto_summary["Total CBM"]), format="%.3f", key=f"sum_cbm_{current_doc_key}")
            total_gw = sc3.number_input("Total GW", min_value=0.0, value=float(auto_summary["Total GW"]), format="%.2f", key=f"sum_gw_{current_doc_key}")
            total_nw = sc4.number_input("Total NW", min_value=0.0, value=float(auto_summary["Total NW"]), format="%.2f", key=f"sum_nw_{current_doc_key}")
        else:
            total_boxes = auto_summary["Total Boxes"]
            total_cbm = auto_summary["Total CBM"]
            total_gw = auto_summary["Total GW"]
            total_nw = auto_summary["Total NW"]
            sc1.metric("Total Boxes", total_boxes)
            sc2.metric("Total CBM", f"{total_cbm:.3f}")
            sc3.metric("Total GW", f"{total_gw:.2f} KG")
            sc4.metric("Total NW", f"{total_nw:.2f} KG")

        packing_summary_values = {
            "Total Boxes": int(total_boxes),
            "Total CBM": float(total_cbm),
            "Total GW": float(total_gw),
            "Total NW": float(total_nw),
            "Manual Override": bool(packing_manual_override),
        }

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.warning("Proforma: packing list is hidden. It will auto-create when converted to Invoice.")

    bill_to = {"Company Name":bill_company,"Registration Number":bill_reg,"GST/VAT":bill_vat,"Contact Person":bill_contact,"Phone":bill_phone,"Email":bill_email,"Country":bill_country,"Address":bill_address}
    docdata = {
        "id": editing.get("id") if editing else str(uuid.uuid4()),
        "type": doc_type, "number": doc_number, "date": str(doc_date), "currency": currency,
        "bill_to": bill_to, "ship_same": ship_same, "ship_to": ship_to,
        "products": products, "discount_type": discount_type, "discount_value": discount_value,
        "shipping_enabled": shipping_enabled, "shipping_cost": shipping_cost,
        "vat_mode": vat_mode, "vat_rate": VAT_RATE, "vat_amount": vat_amount,
        "terms": terms, "packing": packing, "packing_summary": packing_summary_values, "total": grand,
        "created_at": editing.get("created_at") if editing else datetime.now().isoformat(timespec="seconds"),
        "updated_at": datetime.now().isoformat(timespec="seconds")
    }

    x1,x2,x3 = st.columns(3)
    if x1.button("Save / Update Document", type="primary"):
        # save/update customer
        if bill_company:
            cust = dict(bill_to)
            cust["ship_to"] = bill_to if ship_same else ship_to
            found = next((i for i,c in enumerate(customers) if c.get("Company Name","").lower()==bill_company.lower()), None)
            if found is None:
                customers.append(cust)
            else:
                customers[found] = cust
            save_json(CUSTOMERS_FILE, customers)

        # IMPORTANT: if editing existing document, update the same record by ID.
        # If creating new, append only once.
        existing_id = editing.get("id") if editing else docdata.get("id")
        docdata["id"] = existing_id

        idx = next((i for i,d in enumerate(documents) if d.get("id") == existing_id), None)
        if idx is None:
            documents.append(docdata)
        else:
            documents[idx] = docdata

        save_json(DOCUMENTS_FILE, documents)
        st.session_state.pop("imported_word_docdata", None)

        # After save/update, return to Saved Documents list.
        st.session_state.editing_id = None
        st.session_state.force_page = "Saved Documents"
        st.success(f"Saved / Updated: {docdata['number']}")
        st.rerun()

    if x2.button("Convert Proforma to Invoice"):
        if doc_type == "Proforma Invoice":
            newdoc = dict(docdata)
            newdoc["id"] = str(uuid.uuid4())
            newdoc["type"] = "Invoice"
            newdoc["number"] = next_number("Invoice", documents)
            newdoc["packing"] = packing_from_products(products)
            newdoc["created_at"] = datetime.now().isoformat(timespec="seconds")
            newdoc["updated_at"] = datetime.now().isoformat(timespec="seconds")
            documents.append(newdoc)
            save_json(DOCUMENTS_FILE, documents)
            st.session_state.pop("imported_word_docdata", None)
            st.session_state.editing_id = newdoc["id"]
            st.session_state.page = "Create / Edit"
            st.success(f"Converted to Invoice: {newdoc['number']}. Opening converted invoice for editing/packing details...")
            st.rerun()
        else:
            st.info("Already an Invoice.")

    x3.download_button("Download PDF", data=build_pdf(docdata), file_name=f"{doc_number.replace('/','-')}.pdf", mime="application/pdf")
    st.download_button("Download Word", data=build_word(docdata), file_name=f"{doc_number.replace('/','-')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if st.session_state.page == "Saved Documents":
    st.markdown("<div class='card'><h3 class='gold'>Saved Documents — Select / Edit / Convert / Delete</h3>", unsafe_allow_html=True)
    search = st.text_input("Search saved documents by number, customer, type, currency")
    results = documents
    if search:
        s = search.lower()
        results = [d for d in documents if s in json.dumps(d, ensure_ascii=False).lower()]

    if not results:
        st.warning("No documents saved.")
    else:
        st.caption("Each document has its own Edit, Convert, Download and Delete action. Edit/Convert will take you back to the data entry page.")
        for d in sorted(results, key=lambda x: x.get("updated_at", x.get("created_at","")), reverse=True):
            customer_name = d.get("bill_to", {}).get("Company Name", "")
            row = st.container(border=True)
            with row:
                c0, c1, c2, c3, c4, c5 = st.columns([2.4, 1.3, 2.2, 1.1, 1.2, 1.8])
                c0.markdown(f"### {d.get('number','')}")
                c0.caption(f"{d.get('date','')} · {d.get('currency','')}")
                c1.markdown(f"**{d.get('type','')}**")
                c2.markdown(customer_name or "No customer")
                c3.markdown(f"**{money(d.get('total',0), d.get('currency','EUR'))}**")

                if c4.button("Edit", key=f"edit_{d.get('id')}"):
                    st.session_state.editing_id = d.get("id")
                    st.session_state.force_page = "Create / Edit"
                    st.rerun()

                if d.get("type") == "Proforma Invoice":
                    if c5.button("Convert to Invoice", key=f"convert_{d.get('id')}"):
                        newdoc = dict(d)
                        newdoc["id"] = str(uuid.uuid4())
                        newdoc["type"] = "Invoice"
                        newdoc["number"] = next_number("Invoice", documents)
                        newdoc["packing"] = packing_from_products(newdoc.get("products", []))
                        newdoc["created_at"] = datetime.now().isoformat(timespec="seconds")
                        newdoc["updated_at"] = datetime.now().isoformat(timespec="seconds")
                        documents.append(newdoc)
                        save_json(DOCUMENTS_FILE, documents)
                        st.session_state.editing_id = newdoc["id"]
                        st.session_state.force_page = "Create / Edit"
                        st.rerun()
                else:
                    c5.caption("Already invoice")

                d1, d2 = st.columns([1.2, 2])
                d1.download_button(
                    "Download PDF",
                    data=build_pdf(d),
                    file_name=f"{d.get('number','document').replace('/','-')}.pdf",
                    mime="application/pdf",
                    key=f"pdf_{d.get('id')}"
                )
                d1.download_button(
                    "Download Word",
                    data=build_word(d),
                    file_name=f"{d.get('number','document').replace('/','-')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{d.get('id')}"
                )

                confirm = d2.checkbox(f"Confirm delete {d.get('number','')}", key=f"confirm_delete_{d.get('id')}")
                if d2.button("Delete", key=f"delete_{d.get('id')}", type="secondary"):
                    if confirm:
                        documents = [x for x in documents if x.get("id") != d.get("id")]
                        save_json(DOCUMENTS_FILE, documents)
                        if st.session_state.editing_id == d.get("id"):
                            st.session_state.editing_id = None
                        st.session_state.force_page = "Saved Documents"
                        st.success(f"Deleted {d.get('number','document')}")
                        st.rerun()
                    else:
                        st.warning("Tick confirm delete first.")
    st.markdown("</div>", unsafe_allow_html=True)

if st.session_state.page == "Customers":
    st.markdown("<div class='card'><h3 class='gold'>Customer Database</h3>", unsafe_allow_html=True)
    if customers:
        st.dataframe(pd.DataFrame(customers).drop(columns=["ship_to"], errors="ignore"), use_container_width=True, hide_index=True)
        customer_names_for_delete = [c.get("Company Name","") for c in customers if c.get("Company Name","")]
        if customer_names_for_delete:
            selected_customer_delete = st.selectbox("Select customer to delete", customer_names_for_delete)
            confirm_delete_customer = st.checkbox(f"Confirm delete customer {selected_customer_delete}")
            if st.button("Delete Selected Customer", type="secondary"):
                if confirm_delete_customer:
                    customers = [c for c in customers if c.get("Company Name","") != selected_customer_delete]
                    save_json(CUSTOMERS_FILE, customers)
                    st.success(f"Deleted customer: {selected_customer_delete}")
                    st.rerun()
                else:
                    st.warning("Please tick confirm delete first.")
    else:
        st.info("No customers saved yet.")
    st.markdown("</div>", unsafe_allow_html=True)

if st.session_state.page == "Settings":
    st.markdown("<div class='card'><h3 class='gold'>Settings</h3>", unsafe_allow_html=True)
    st.write("Password is currently set to **1985**.")
    new_terms = st.text_area("Default Terms & Conditions", value=settings.get("terms",""), height=430)
    if st.button("Save Default Terms"):
        settings["terms"] = new_terms
        save_json(SETTINGS_FILE, settings)
        st.success("Settings saved.")
    st.markdown("</div>", unsafe_allow_html=True)
